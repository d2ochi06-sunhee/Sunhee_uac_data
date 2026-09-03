# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def apply_font(run, font_name='맑은 고딕', size_pt=10.5, bold=False, color_rgb=None):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color_rgb:
        run.font.color.rgb = color_rgb
    
    # Set East Asia font for Word CJK rendering
    rPr = run._element.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:eastAsia="{font_name}"/>')
    rPr.append(rFonts)

doc = Document()

# Default normal font
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10.5)

# Page margins
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Styles & Colors
NAVY = RGBColor(26, 54, 93)
ACCENT_BLUE = RGBColor(49, 130, 206)
DARK_GRAY = RGBColor(45, 55, 72)
ORANGE = RGBColor(221, 107, 32)
GRAY = RGBColor(113, 128, 150)

# Title
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_title = title_p.add_run('🎤 [발표 대본] 외래관광객조사 데이터 기반\nK-컬처 영향 분석 및 2026 관광 전략')
apply_font(run_title, '맑은 고딕', 20, True, NAVY)

# Subtitle / Info
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_sub = sub_p.add_run('기반 자료: kculture_deck_데이터정합성_보강본.pptx (총 13개 슬라이드) | 소요시간: 약 5~7분')
apply_font(run_sub, '맑은 고딕', 10, False, GRAY)

doc.add_paragraph() # Spacer

slides_data = [
    {
        'num': 1,
        'title': '타이틀 (Title)',
        'screen': '메인 제목 "TOURISM ANALYTICS: K-컬처 영향 외래관광객 체재 및 지출 특성 분석"과 핵심 수치 요약(N=30,347, -5.8일, 79.0%) 화면',
        'action': '청중을 부드럽게 둘러보며 인사를 건네고, 명확하고 또박또박한 톤으로 발표를 시작합니다.',
        'script': '“안녕하십니까, 『2023년~2025년 외래관광객조사 데이터 기반 K-컬처 영향 분석 및 2026 전략 수립』 발표를 맡은 [이름]입니다.\n\n최근 몇 년 사이 명동이나 홍대, 성수동을 가보면 외국인 관광객들이 정말 많이 늘었다는 걸 체감하실 겁니다. 그런데 이분들이 과연 단순히 Sightseeing, 즉 풍경만 보러 올까요? 아니면 K-Pop이나 K-뷰티 같은 한류 문화를 즐기러 올까요?\n\n저희 팀은 지난 3년간 한국을 찾은 여가 관광객 3만여 명(30,347명)의 원시 데이터를 정밀하게 파헤쳐, \'K-컬처가 외래관광객의 여행 행태에 미친 실제 영향\'을 분석했습니다.\n\n오늘 이 자리를 통해 데이터가 말해주는 핵심 인사이트와 앞으로 우리가 나아가야 할 2026년 관광 전략을 알기 쉽게 풀어드리겠습니다.”',
        'checkpoint': '최근 3개년 3만 명 데이터를 파헤쳐 K-컬처 관광객의 진짜 특징과 2026년 전략을 발표하겠다.'
    },
    {
        'num': 2,
        'title': '연구 배경 & 목적 (Background & Objectives)',
        'screen': '한류/K-Pop, K-뷰티 이미지와 함께 연구 필요성 및 목표 3가지가 표시된 화면',
        'action': '손짓으로 화면을 가리키며 관광 패러다임 변화를 강조합니다.',
        'script': '“먼저 2번 슬라이드, 연구 배경과 목적입니다.\n\n과거의 한국 관광은 \'경복궁 가기\', \'N타워 올라가기\' 같은 단순 유적지 탐방이 주를 이뤘습니다. 하지만 최근에는 K-Pop 콘서트 관람, 성수동 로드숍 쇼핑, K-뷰티 체험, 한국 맛집 탐방처럼 \'체험형 K-컬처 관광\'으로 패러다임이 완전히 바뀌었습니다.\n\n따라서 저희 연구의 목적은 크게 3가지입니다.\n첫째, K-컬처에 영향을 받은 관광객이 일반 관광객과 비교했을 때 머무는 기간(체재일수)과 돈을 쓰는 방식(지출)이 어떻게 다른가?\n둘째, 연도별로 이들의 비중이 얼마나 늘고 있는가?\n셋째, 이 데이터를 바탕으로 2026년에 관광객들을 끌어모을 타깃 마케팅 전략을 어떻게 세울 것인가입니다.”',
        'checkpoint': 'K-컬처 영향으로 변화된 외래관광객의 체재·지출 특성을 파악하고 타깃 전략을 세우는 것이 목적이다.'
    },
    {
        'num': 3,
        'title': '분석 데이터 개요 (Data Overview)',
        'screen': '2023, 2024, 2025 데이터셋 수치(총 48,522명 ➔ 표본 30,347명) 및 파이프라인 구조도',
        'action': '신뢰감을 줄 수 있도록 차분하고 신뢰성 있는 목소리로 수치를 정확히 전달합니다.',
        'script': '“3번 슬라이드는 저희가 사용한 데이터의 개요입니다.\n\n저희는 문화체육관광부와 한국문화관광연구원에서 제공하는 2023년, 2024년, 2025년 외래관광객조사 데이터를 통합 활용했습니다.\n\n전체 조사 응답자 48,522명 중에서, 비즈니스나 학업 목적이 아닌 순수 여가·휴양 목적으로 방한한 관광객(D_MOK=1) 30,347명을 엄선하여 분석 대상(N)으로 삼았습니다.\n\n또한 통계적 정확도를 높이기 위해 정부 통계 표본 가중치(weight)를 엄격히 적용하였으며, 17개 시도별 방문 일수와 지출 항목 파생변수까지 완벽하게 정제하여 분석을 진행했습니다.”',
        'checkpoint': '2023~2025년 3개년 여가 관광객 30,347명의 정부 원시 데이터를 표본가중치까지 반영해 분석했다.'
    },
    {
        'num': 4,
        'title': '타겟 그룹 정의 (Group Definition)',
        'screen': '\'K-컬처 관여군(한류관광객)\' vs \'일반 관광객\' 그룹 구분 기준 및 비율 그래프',
        'action': '양손으로 두 그룹을 대비시키는 동작을 취합니다.',
        'script': '“그렇다면 저희는 \'K-컬처 관광객\'을 어떻게 정의했을까요? 4번 슬라이드를 봐주시기 바랍니다.\n\n한국에 올 때 방문 고려 요인(Q1_1, Q2) 중 K-Pop, K-드라마, K-뷰티, 한류 스타/굿즈, 쇼핑 등 한류 관련 항목을 1개 이상 선택한 관광객을 \'K-컬처 관여군\'으로 분류했습니다. 그리고 이에 해당하지 않는 분들을 \'일반 관광객\'으로 나눠 비교 분석했습니다.\n\n놀라운 점은, 여가 관광객 중 K-컬처 관여군의 비율이 초기 41.0%에서 최근 49.8%까지 지속적으로 상승하며 이제 전체 관광객의 절반에 육박하고 있다는 사실입니다.”',
        'checkpoint': '한류/K-뷰티/쇼핑을 고려해 온 \'K-컬처 관여군\'이 전체 여가 관광객의 절반(49.8%)에 달한다.'
    },
    {
        'num': 5,
        'title': '주요 발견 1 - 체재 기간 분석 (Stay Duration)',
        'screen': '체재일수 갭(-5.8일, 34.6일 vs 40.4일) 및 국적별/연령별 그래프',
        'action': '핵심 반전 수치인 \'-5.8일\'과 \'48%\'를 말할 때 목소리 톤을 높입니다.',
        'script': '“자, 이제 가장 중요한 주요 발견(Key Findings)입니다. 5번 슬라이드를 주목해 주십시오.\n\n첫 번째 반전은, K-컬처 관광객의 체재 기간이 일반 관광객보다 평균 \'5.8일\' 짧다는 점입니다. (K-컬처 관여군 34.6일 vs 일반군 40.4일)\n\n\'어? 한류를 좋아하면 더 오래 머물지 않나?\'라고 생각하실 수 있는데요, 그 이유는 두 가지입니다.\n\n첫째, K-컬처 팬의 절반 가까이(48.0%)가 15세~29세의 청년층입니다. 일반 관광객(28.7%)보다 젊은 층 비중이 19.3%p나 높습니다.\n둘째, 이들은 일본(-7.6일), 대만(-6.2일), 중국(-5.4일)처럼 가까운 아시아 국가에서 방학이나 휴가를 이용해 \'짧고 알차게\' 다녀가는 경향이 강하기 때문입니다.”',
        'checkpoint': 'K-컬처 팬은 1020 청년층(48%)이 주를 이루며, 일본/대만 등 근거리에서 짧고 알차게(-5.8일) 다녀간다.'
    },
    {
        'num': 6,
        'title': '주요 발견 2 - 지출 및 침투율 분석 (Spend & Trends)',
        'screen': '3개년 K-컬처 관여 비중(36.3% ➔ 42.5%) 및 1인당 지출액 비교 차트',
        'action': '일일 소비 지출액의 의미를 강조하며 고개를 끄덕입니다.',
        'script': '“이어서 6번 슬라이드, 지출과 연도별 침투율 추이입니다.\n\n전체 관광객 중 K-컬처 관여군의 비중은 2023년 36.3% ➔ 2024년 42.2% ➔ 2025년 42.5%로 매년 가파르게 성장하고 있습니다. 15~29세 청년층 비중 역시 48% 수준을 꾸준히 유지하고 있습니다.\n\n1인당 총 지출액을 보면 K-컬처 관여군이 37.9만 원, 일반군이 38.2만 원으로 약 0.3만 원 차이밖에 나지 않습니다. 체재 일수는 6일이나 짧은데 지출 총액은 비슷하다는 뜻은, 하루 단위로 환산했을 때 K-컬처 관광객이 한국에서 훨씬 더 많은 돈을 아낌없이 쓰고 간다는 것을 의미합니다.”',
        'checkpoint': 'K-컬처 비중은 42.5%로 가파르게 상승 중이며, 머무는 기간 대비 일일 소비 지출액이 매우 높다.'
    },
    {
        'num': 7,
        'title': '주요 발견 3 - 활동 및 쇼핑 특성 (Activities & Shopping)',
        'screen': '주요 활동 참여율 79.0% vs 62.7% (+16.3%p) 비교 및 품목별 인포그래픽',
        'action': '\'입고, 바르고, 먹고, 소장하는\' 문구를 또박또박 강조합니다.',
        'script': '“그렇다면 이들은 한국에서 무슨 돈을 그렇게 많이 쓸까요? 7번 슬라이드 활동 특성입니다.\n\nK-컬처 관여 관광객의 79.0%가 쇼핑, K-뷰티, 미식(식도락) 활동에 집중적으로 참여했습니다. 이는 일반 관광객(62.7%)보다 무려 16.3%p나 높은 수치입니다.\n\n특히 구매 품목을 보면 화장품/향수 구매율이 일반 관광객보다 16.4%p 높았고, 패션 의류, K-Pop 굿즈/앨범 구매에 지출이 크게 집중되어 있었습니다. 이들은 한국 문화를 단순히 보는 것에 그치지 않고 \'입고, 바르고, 먹고, 소장하는\' 적극적 소비층입니다.”',
        'checkpoint': 'K-컬처 관광객의 79%가 쇼핑·뷰티·미식에 집중하며, 화장품/의류/굿즈 소비에 적극적이다.'
    },
    {
        'num': 8,
        'title': '주요 발견 4 - 만족도 및 충성도 (Satisfaction & Loyalty)',
        'screen': '만족도(4.63점 vs 4.50점), 재방문/추천 의향 비교 막대그래프',
        'action': '충성도 높은 결과에 확신에 찬 표정을 지어 보입니다.',
        'script': '“8번 슬라이드는 만족도와 충성도 분석입니다.\n\n한국 여행 전반에 대한 만족도를 5점 만점으로 조사한 결과, K-컬처 관여군은 4.63점으로 일반 관광객(4.50점)보다 훨씬 높은 만족도를 보였습니다.\n\n덕욱 놀라운 것은 충성도 지표입니다. 타인 추천 의향은 4.67점, 재방문 의향 역시 4.63점으로 매우 높게 나타났습니다.\n자신이 좋아하는 문화(K-Pop, K-뷰티 등)를 현지에서 직접 체험하고 돌아간 관광객일수록 한국에 대한 호감도가 극대화되고, 주변에 적극 추천하며, 반드시 다시 한국을 찾는다는 사실이 증명된 것입니다.”',
        'checkpoint': 'K-컬처 관광객은 전반적 만족도(4.63점)와 재방문/추천 의향(4.67점)이 매우 높은 \'충성 고객\'이다.'
    },
    {
        'num': 9,
        'title': '종합 인사이트 (Comprehensive Insights)',
        'screen': '3대 핵심 특징 요약 카드 및 챌린지 요약 (짧은 체재기간 보완 필요성)',
        'action': '손가락으로 1, 2, 3을 세며 핵심 포인트를 정리합니다.',
        'script': '“9번 슬라이드, 지금까지의 데이터를 3가지로 종합해 보겠습니다.\n\n첫째, 비중의 확대: K-컬처 영향 관광객은 36%에서 42.5%로 계속 늘고 있습니다.\n둘째, 타깃의 명확성: 15~29세 청년층이 48%이며, 쇼핑·뷰티·미식 참여율이 79%에 달합니다.\n셋째, 높은 충성도와 명확한 과제: 만족도와 재방문 의향은 최고 수준이지만, 체재 일수가 5.8일 짧다는 확실한 약점이 존재합니다.\n\n결국 우리의 핵심 과제는 \'이 짧게 머무는 젊은 K-컬처 팬들의 체재 기간을 어떻게 늘리고, 머무는 동안 소비를 어떻게 극대화할 것인가\'에 있습니다.”',
        'checkpoint': '핵심 과제는 \'젊고 충성도 높은 K-컬처 팬들의 짧은 체재 기간을 늘리고 소비를 극대화하는 것\'이다.'
    },
    {
        'num': 10,
        'title': '전략적 제언 - 3대 액션 플랜 (Strategic Action Items)',
        'screen': 'Action 1, Action 2, Action 3 카드 3개로 구성된 전략 제언 화면',
        'action': '목소리에 힘을 주어 해결책을 자신감 있게 발표합니다.',
        'script': '“이에 저희 팀은 2026년을 향한 3대 전략적 액션 플랜(Action Items)을 제언합니다. 10번 슬라이드입니다.\n\nACTION 1: 일본/대만 15-29세 타깃 \'2박 3일 초밀도 K-체험 패키지\' 개발\n체재일수가 가장 짧은 일본(-7.6일)과 대만(-6.2일) 청년층을 위해, 성수동 로드숍, K-Pop 안무 클래스, 인스타 핫플을 묶은 2박 3일 고밀도 코스를 제공해 단기 체재의 한계를 극복해야 합니다.\n\nACTION 2: 중국/동남아 2030 타깃 \'K-뷰티 + 미식 융합 프리미엄 상품\'\n피부과/메이크업 체험과 미쉐린 K-푸드 투어를 결합하여 1인당 지출 단가를 대폭 끌어올려야 합니다.\n\nACTION 3: LCC(저비용항공사) 연계 A/B 결합 패키지\n개별여행(FIT) 비중이 높은 젊은 층을 위해, 항공권과 공연 티켓, 시내면세점 할인 쿠폰을 하나로 묶은 통합 바우처를 제공하는 것입니다.”',
        'checkpoint': '① 일본/대만 2박3일 밀도높은 K-체험, ② 중국/동남아 뷰티+미식 융합, ③ LCC 연계 통합 바우처 제공.'
    },
    {
        'num': 11,
        'title': '대시보드 시스템 구축 (Dashboard System)',
        'screen': '구축된 K-컬처 데이터 대시보드(index.html) UI 캡처 화면',
        'action': '대시보드가 실용적임을 강조하며 미소를 띱니다.',
        'script': '“11번 슬라이드는 저희가 실제로 구현한 \'K-컬처 관광 실시간 분석 대시보드\' 시스템입니다.\n\n앞서 말씀드린 3만 명의 데이터를 바탕으로, 연도별, 국적별, 연령대별로 K-컬처 관여율, 체재일수 갭, 지출 항목을 클릭 한 번으로 실시간 비교·분석할 수 있는 웹 대시보드를 구축했습니다.\n\n이를 통해 마케팅 담당자들은 2026년 신규 수립될 관광 정책의 성과를 데이터 기반으로 실시간 모니터링할 수 있게 됩니다.”',
        'checkpoint': '실시간으로 국적·연령별 K-컬처 관광 특성을 분석할 수 있는 웹 대시보드를 구축 완료했다.'
    },
    {
        'num': 12,
        'title': '데이터 처리 및 분석 방법론 (Data Pipeline)',
        'screen': 'analyze.py ➔ CSV/JS ➔ 대시보드/PPT로 이어지는 데이터 정합성 파이프라인',
        'action': '차분하고 전문성 있는 어조로 데이터 정밀성을 설명합니다.',
        'script': '“12번 슬라이드는 데이터 정합성 검증 절차입니다.\n\n저희는 2023, 2024, 2025년 엑셀 원시 데이터부터 출발하여 파이썬(analyze.py) 파이프라인을 구축해 수치 검증을 완료했습니다.\n\n데이터의 오류나 누락 없이 정확한 수치만을 선별하여 대시보드와 이번 발표 자료에 반영했음을 말씀드립니다.”',
        'checkpoint': '파이썬 파이프라인을 통해 원시 데이터 수치 정합성을 100% 검증하여 반영했다.'
    },
    {
        'num': 13,
        'title': 'Q&A 및 마무리 (Conclusion & Q&A)',
        'screen': '핵심 메시지 요약 및 "Q&A / 감사합니다" 문구',
        'action': '청중을 향해 허리를 숙여 깍듯이 인사합니다.',
        'script': '“마지막 13번 슬라이드입니다.\n\n오늘 발표의 핵심 메시지는 단 하나입니다.\nK-컬처 외래관광객은 \'젊고, 짧게 머물지만, 강력하게 소비하며, 한국을 반드시 다시 찾는 최고의 고객\'이라는 점입니다.\n\n이들의 특성에 맞춘 고밀도 K-체험 상품과 뷰티·미식 융합 전략이 준비된다면, 2026년 대한민국 관광 산업은 한 단계 더 도약할 것입니다.\n\n이상으로 발표를 마치겠습니다. 경청해 주셔서 감사합니다. 질문이 있으시면 답변해 드리겠습니다.”',
        'checkpoint': 'K-컬처 관광객 맞춤형 전략으로 2026년 한국 관광 산업의 도약을 이끌겠다며 마무리.'
    }
]

for s in slides_data:
    # Section Header
    h = doc.add_heading(f'슬라이드 {s["num"]}. {s["title"]}', level=1)
    for run in h.runs:
        apply_font(run, '맑은 고딕', 14, True, NAVY)
    
    # Info Table Box (Screen & Action)
    tbl = doc.add_table(rows=2, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    tbl.rows[0].cells[0].text = '화면 구성'
    tbl.rows[0].cells[1].text = s['screen']
    tbl.rows[1].cells[0].text = '발표자 행동 팁'
    tbl.rows[1].cells[1].text = s['action']
    
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    apply_font(r, '맑은 고딕', 9.5)
    
    # Script Title
    sp = doc.add_paragraph()
    sp_run = sp.add_run('💬 발표자 구어체 스크립트:')
    apply_font(sp_run, '맑은 고딕', 11, True, ACCENT_BLUE)
    
    # Script Box / Paragraph
    script_p = doc.add_paragraph()
    script_p.paragraph_format.left_indent = Inches(0.2)
    s_run = script_p.add_run(s['script'])
    apply_font(s_run, '맑은 고딕', 10.5, False, DARK_GRAY)
    
    # Checkpoint Box
    cp_p = doc.add_paragraph()
    cp_p.paragraph_format.left_indent = Inches(0.2)
    cp_run1 = cp_p.add_run('💡 1줄 체크포인트: ')
    apply_font(cp_run1, '맑은 고딕', 9.5, True, ORANGE)
    
    cp_run2 = cp_p.add_run(s['checkpoint'])
    apply_font(cp_run2, '맑은 고딕', 9.5, False, DARK_GRAY)
    
    doc.add_paragraph() # Spacer

# Footer Tips Section
h_footer = doc.add_heading('📌 발표자 꿀팁 및 질의응답 대비 요약', level=1)
for run in h_footer.runs:
    apply_font(run, '맑은 고딕', 14, True, NAVY)

tip1 = doc.add_paragraph()
t1 = tip1.add_run('1. 전달력 높은 수치 말하기: ')
apply_font(t1, '맑은 고딕', 10.5, True)
t1_text = tip1.add_run('"48%" 대신 "절반에 가까운 48%", "-5.8일" 대신 "약 6일 정도 짧은 -5.8일"처럼 직관적 비교 문구를 덧붙여 발표하세요.')
apply_font(t1_text, '맑은 고딕', 10.5)

tip2 = doc.add_paragraph()
t2 = tip2.add_run('2. 질의응답 대비 3줄 요약:\n')
apply_font(t2, '맑은 고딕', 10.5, True)
t2_text = tip2.add_run('  • 누가 오나요? 10~20대 청년층 K-컬처 팬 (48%)\n  • 어떻게 여행하나요? 일본/대만 등 근거리에서 짧고 알차게(-5.8일) 와서 쇼핑·뷰티·식도락(79%) 집중 소비\n  • 전략은 무엇인가요? 2박3일 고밀도 K-체험 패키지 및 뷰티+미식 융합 상품 개발')
apply_font(t2_text, '맑은 고딕', 10.5)

paths = [
    r'c:\Users\d2och\Downloads\실습 netflix\2019~25 외래관광객조사_Data\K컬처_외래관광객_분석_발표대본.docx',
    r'c:\Users\d2och\Downloads\실습 netflix\2019~25 외래관광객조사_Data\K_Culture_Survey_Presentation_Script.docx'
]

for p in paths:
    doc.save(p)

print('Successfully saved Word documents with UTF-8 and EastAsia Korean Font settings.')
